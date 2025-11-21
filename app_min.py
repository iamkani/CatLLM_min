
import os
import io
import math
import re
import json, gzip
from pathlib import Path
from collections import Counter
from typing import List, Dict, Tuple

import streamlit as st
import pandas as pd

from prompts import get_fewshots_for_role

# Configure the page early before any UI is drawn
st.set_page_config(page_title="CatLLM — Minimal (Docs)", page_icon="🐾", layout="wide")

FEWSHOT_ENABLED = os.getenv("FEWSHOT_ENABLED", "1") not in {"0", "false", "False"}
FEWSHOT_PAIRS = int(os.getenv("FEWSHOT_PAIRS", "2"))

SOURCE_MAP = {
    "USDA AMS": "https://www.ams.usda.gov/rules-regulations/research-promotion/beef",
    "USDA": "https://www.ams.usda.gov/",
    "GAO": "https://www.gao.gov/assets/gao-18-55r.pdf",
    "FSIS": "https://www.fsis.usda.gov/inspection/compliance-guidance/labeling/claims-guidance",
    "FoodSafety.gov": "https://www.foodsafety.gov/food-safety-charts/safe-minimum-internal-temperatures",
    "AMS Grades": "https://www.ams.usda.gov/grades-standards/carcass-beef-grades-and-standards",
    "NASS": "https://www.nass.usda.gov/Quick_Stats/",
    "ERS": "https://www.ers.usda.gov/topics/animal-products/cattle-beef/sector-at-a-glance/",
}

# ---------- linkify helpers ----------
_BRACKETED = re.compile(r"\[([^\[\]]+?)\](?!\()", re.UNICODE)
_LABELS = sorted(SOURCE_MAP.keys(), key=len, reverse=True)
# Build the alternation of known labels (already sorted by length above)
_ALT = "|".join(re.escape(k) for k in _LABELS)

_BARE = re.compile(
    rf"(?<!\[)(?:(?<=^)|(?<=[^\w-]))({_ALT})(?:(?=$)|(?=[^\w-]))",
    re.IGNORECASE | re.UNICODE,
)

def _lookup(urlmap, key: str) -> str | None:
    # case-insensitive lookup
    for k, v in urlmap.items():
        if k.lower() == key.lower():
            return v
    return None

def linkify_labels(text: str) -> str:
    def _lookup(urlmap, key: str) -> str | None:
        for k, v in urlmap.items():
            if k.lower() == key.lower():
                return v
        return None

    # 1) Linkify things already in square brackets like: [USDA AMS], [GAO], ...
    def _sub_bracketed(m):
        inside = m.group(1).strip()
        norm = inside.replace("\u2013", "-").replace("\u2014", "-")
        base = re.split(r"[|\-:]", norm, maxsplit=1)[0].strip()
        url = _lookup(SOURCE_MAP, base)
        return f"[{inside}]({url})" if url else m.group(0)

    out = _BRACKETED.sub(_sub_bracketed, text)

    # 2) Linkify bare labels only when they’re standalone words/phrases
    def _sub_bare(m):
        label = m.group(1)
        url = _lookup(SOURCE_MAP, label)
        return f"[{label}]({url})" if url else label

    out = _BARE.sub(_sub_bare, out)
    return out

# ---------- loader for prebuilt corpora (.json or .json.gz) ----------
@st.cache_data(show_spinner=False)
def load_json_list(path: str) -> list[str]:
    """Load list[str] from .json or .json.gz. Returns [] if not found or invalid."""
    p = Path(path)
    if not p.exists():
        gz = p.with_suffix(p.suffix + ".gz") if p.suffix != ".gz" else p
        if gz.exists():
            p = gz
        else:
            return []
    try:
        if p.suffix.endswith("gz"):
            with gzip.open(p, "rt", encoding="utf-8") as f:
                return json.load(f)
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception as e:
        st.warning(f"Failed to load {p.name}: {e}")
        return []

# ------------------------
# Role definitions & credentials
# ------------------------
ROLE_SYSTEM_MESSAGES: Dict[str, str] = {
    "Association Analyst": (
        "You are an Association Analyst. "
        "You specialize in interpreting association records, market trends, and industry benchmarks for cattle operations. "
        "Provide data-driven insights and analyses, and quote sources when relevant. "
        "Prefer grounded answers with inline quotes like: > snippet [source]. If unsure, say so."
    ),
    "Buyer / Feeder": (
        "You are a Buyer / Feeder. "
        "You focus on purchasing and feeding cattle, and provide guidance on feed strategies, pricing, and buying decisions. "
        "Quote sources when relevant and prefer grounded answers with inline quotes like: > snippet [source]. If unsure, say so."
    ),
    "Genetic Advisor": (
        "You are a Genetic Advisor. "
        "You specialize in cattle genetics, breeding strategies, and selection. "
        "Provide advice grounded in genetics and science, quoting sources when relevant. "
        "Prefer grounded answers with inline quotes like: > snippet [source]. If unsure, say so."
    ),
    "Independent Rancher": (
        "You are an Independent Rancher. "
        "You run a cattle ranch independently, focusing on raising cattle, pasture management, and general operations. "
        "Provide practical guidance and insights, quoting sources when relevant. "
        "Prefer grounded answers with inline quotes like: > snippet [source]. If unsure, say so."
    ),
    "Root": (
        "You are the root user with full access. "
        "Provide concise, helpful answers with inline quotes like: > snippet [source]. "
        "If unsure, say so."
    ),
}

ROLE_ALIASES = {
    "association_analyst": "Association Analyst",
    "buyer_feeder": "Buyer / Feeder",
    "genetic_advisor": "Genetic Advisor",
    "independent_rancher": "Independent Rancher",
    "root": "Root",
}

def normalize_role(role: str | None) -> str:
    if not role:
        return "Root"
    r = role.strip().lower().replace("/", "_").replace(" ", "_")
    return ROLE_ALIASES.get(r, ROLE_ALIASES.get("association_analyst"))

def parse_credentials() -> tuple[dict[str, str], dict[str, tuple[str, str]]]:
    """
    Parse credentials from env:
      CATLLM_ROOT_CREDS = "admin:supersecret,ops:opspass"      (pairs)
      CATLLM_USERS      = "alexa:pw1:association_analyst,..."  (triplets)
    Returns:
      roots:  {username: password}
      users:  {username: (password, normalized_role)}
    """
    roots: Dict[str, str] = {}
    for part in filter(None, [p.strip() for p in os.getenv("CATLLM_ROOT_CREDS", "").split(",")]):
        if ":" in part:
            u, pw = part.split(":", 1)
            roots[u.strip()] = pw.strip()

    users: Dict[str, tuple[str, str]] = {}
    for part in filter(None, [p.strip() for p in os.getenv("CATLLM_USERS", "").split(",")]):
        bits = part.split(":")
        if len(bits) != 3:
            continue
        u, pw, role = bits[0].strip(), bits[1].strip(), bits[2].strip()
        users[u] = (pw, normalize_role(role))

    if not roots and not users:
        roots = {"root": "root123"}
        users = {"user": ("user123", normalize_role("association_analyst"))}

    return roots, users

# ------------------------
# Optional deps detection
# ------------------------
def has_pkg(name: str) -> bool:
    try:
        __import__(name)
        return True
    except Exception:
        return False

HAS_OPENAI = has_pkg("openai")
HAS_PYPDF = has_pkg("pypdf")
HAS_DOCX = has_pkg("docx")  # python-docx
HAS_OPENPYXL = has_pkg("openpyxl")
HAS_XLRD = has_pkg("xlrd")

# ------------------------
# Minimal model call
# ------------------------
def minimal_model_reply(
    user_text: str,
    history: List[Dict],
    context_chunks: List[str],
    role_name: str = None,
) -> str:
    # Build a preface that embeds retrieved context (if any)
    ctx_joined = "\n\n---\n".join(context_chunks[:4])
    preface = (
        "Use the following context to answer naturally. "
        "When referencing information, explicitly mention the document name and, if available, the page number. "
        "For example: 'In Feb2022HW.pdf, page 39, you can find...' "
        "Avoid saying 'the documents' or 'the provided materials'—use the actual file names provided below. "
        "Be clear and conversational, but make sure the reader can trace each fact to a source file.\n\n"
        + ctx_joined
        + "\n\n"
    )

    # Determine system prompt based on role
    role_canonical = normalize_role(role_name)
    system_prompt = ROLE_SYSTEM_MESSAGES.get(role_canonical, ROLE_SYSTEM_MESSAGES["Root"])

    # Few-shots (role-grounded)
    fewshots = get_fewshots_for_role(role_canonical) if FEWSHOT_ENABLED else []
    max_msgs = max(0, 2 * FEWSHOT_PAIRS)   # keep full user/assistant pairs
    fewshots = fewshots[:max_msgs]
    if len(fewshots) % 2 == 1:
        fewshots = fewshots[:-1]

    try:
        from openai import OpenAI  # type: ignore
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError("OPENAI_API_KEY not set")
        client = OpenAI(api_key=api_key)

        # Build the message list (system → few-shots → recent history → current user)
        msgs = [{"role": "system", "content": system_prompt}]
        msgs += fewshots

        for m in history[-6:]:
            r = m.get("role", "user")
            c = m.get("content", "")
            if r in {"user", "assistant"} and c:
                msgs.append({"role": r, "content": c})

        msgs.append({"role": "user", "content": preface + user_text})

        resp = client.chat.completions.create(
            model=os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini"),
            messages=msgs,
            temperature=0.2,
        )
        return resp.choices[0].message.content or "(no content)"

    except Exception:
        # Local fallback: surface best snippets
        if context_chunks:
            snippets = "\n\n---\n".join(context_chunks[:3])
            return (
                "(local heuristic)\nTop snippets:\n\n"
                + snippets
                + "\n\nYour question: "
                + user_text
            )
        return f"(minimal echo) {user_text}"

# ------------------------
# Simple text utils
# ------------------------
_STOPWORDS = set(
    (
        "a an and the of to in is it that this for on with as at by from be are was were or if "
        "then so such via into up out over under within without about between across not no yes you "
        "your yours we us our they their i me my mine he she him her his hers its who whom which what "
        "when where why how been being do does did done can could should would may might will shall "
        "just only also etc"
    ).split()
)

def normalize_text(s: str) -> str:
    s = s.replace("\x00", "")
    s = re.sub(r"\s+", " ", s)
    return s.strip()

def tokenize(s: str) -> List[str]:
    return [t for t in re.findall(r"[A-Za-z0-9_]+", s.lower()) if t not in _STOPWORDS]

def chunk_text(s: str, chunk_chars: int = 1200, overlap: int = 150) -> List[str]:
    s = normalize_text(s)
    if len(s) <= chunk_chars:
        return [s]
    chunks: List[str] = []
    start = 0
    while start < len(s):
        end = min(len(s), start + chunk_chars)
        chunks.append(s[start:end])
        if end == len(s):
            break
        start = max(0, end - overlap)
    return chunks

def score_chunk(query: str, chunk: str) -> float:
    # Simple bag-of-words tf scoring
    q_tokens = tokenize(query)
    if not q_tokens:
        return 0.0
    c_counts = Counter(tokenize(chunk))
    return sum(c_counts.get(t, 0) for t in set(q_tokens)) / (1 + math.log10(1 + len(chunk)))

def top_chunks(query: str, chunks: List[str], k: int = 5) -> List[str]:
    scored: List[Tuple[float, str]] = [(score_chunk(query, c), c) for c in chunks]
    scored.sort(key=lambda x: x[0], reverse=True)
    return [c for s, c in scored if s > 0][:k]

# ------------------------
# Extraction for uploads + source-aware chunking
# ------------------------
def extract_text_from_upload(file) -> Tuple[str, str, Dict]:
    """Return (text, info_str, meta) for an uploaded file.

    The ``meta`` dictionary may include 'type', 'pages' (list of page texts), or 'df' for
    tabular data. The ``info_str`` gives a short description of the file. Unsupported
    formats return an empty string and a descriptive message.
    """
    name = file.name
    suffix = name.split(".")[-1].lower()

    try:
        if suffix in {"txt", "md", "log"}:
            data = file.read().decode("utf-8", errors="ignore")
            return normalize_text(data), f"{name} (plain text)", {"type": "text"}

        if suffix in {"csv"}:
            df = pd.read_csv(file)
            preview = df.to_csv(index=False)
            return normalize_text(preview), f"{name} (CSV {df.shape[0]} rows × {df.shape[1]} cols)", {"type": "csv", "df": df}

        if suffix in {"xlsx", "xls"}:
            try:
                df = pd.read_excel(file)  # engine auto
                preview = df.to_csv(index=False)
                return normalize_text(preview), f"{name} (Excel {df.shape[0]} rows × {df.shape[1]} cols)", {"type": "excel", "df": df}
            except Exception as e:
                return "", f"{name} (Excel) could not be parsed: {e}. Install openpyxl/xlrd.", {"type": "excel_error"}

        if suffix in {"docx"}:
            if not HAS_DOCX:
                return "", f"{name} (DOCX) requires 'python-docx' to extract text.", {"type": "docx_error"}
            from docx import Document  # type: ignore
            bio = io.BytesIO(file.read())
            doc = Document(bio)
            text = "\n".join(p.text for p in doc.paragraphs)
            return normalize_text(text), f"{name} (DOCX paragraphs: {len(doc.paragraphs)})", {"type": "docx"}

        if suffix in {"pdf"}:
            if not HAS_PYPDF:
                return "", f"{name} (PDF) requires 'pypdf' to extract text.", {"type": "pdf_error"}
            from pypdf import PdfReader  # type: ignore
            bio = io.BytesIO(file.read())
            reader = PdfReader(bio)
            pages: List[str] = []
            for page in reader.pages:
                try:
                    pages.append(page.extract_text() or "")
                except Exception:
                    pages.append("")
            text = "\n".join(pages)
            return normalize_text(text), f"{name} (PDF pages: {len(reader.pages)})", {"type": "pdf", "pages": pages}

        return "", f"{name}: unsupported file type '{suffix}'.", {"type": "unsupported"}

    except Exception as e:
        return "", f"{name}: error extracting text: {e}", {"type": "error"}


def build_source_marked_chunks(name: str, meta: Dict, text: str) -> List[str]:
    """Create source-marked text chunks for retrieval.

    For PDFs, chunk by page and add page numbers; for CSV/Excel files, chunk
    by row windows; for other types, normal text chunking with a filename marker.
    """
    t = meta.get("type", "text")
    # PDF: chunk by page for clear markers
    if t == "pdf" and meta.get("pages") is not None:
        out: List[str] = []
        for i, pg in enumerate(meta["pages"], start=1):
            for c in chunk_text(pg, chunk_chars=1200, overlap=120):
                out.append(f"[{name} | pdf | page {i}] \n{c}")
        return out
    # CSV/Excel: chunk by row windows for clear row ranges
    if t in {"csv", "excel"} and meta.get("df") is not None:
        df: pd.DataFrame = meta["df"]
        window = 40
        out: List[str] = []
        for start in range(0, len(df), window):
            end = min(len(df), start + window)
            block_csv = df.iloc[start:end].to_csv(index=False)
            for c in chunk_text(block_csv, chunk_chars=1600, overlap=0):
                out.append(f"[{name} | {t} | rows {start+1}–{end}] \n{c}")
        return out
    # DOCX/TXT/others: normal chunking with filename marker
    return [f"[{name} | {t}] \n{c}" for c in chunk_text(text, chunk_chars=1400, overlap=200)]

# ------------------------
# Session state and authentication
# ------------------------
if "messages" not in st.session_state:
    st.session_state.messages = []
if "docs" not in st.session_state:
    st.session_state.docs = []
if "all_chunks" not in st.session_state:
    st.session_state.all_chunks = []
if "auth" not in st.session_state:
    st.session_state.auth = {"is_authed": False, "username": None, "role": None}

# Load prebuilt corpora (priority: rag_store > rag_store_cluster3)
if "stores" not in st.session_state:
    primary = load_json_list("rag_store.json") or load_json_list("rag_store.json.gz")
    secondary = load_json_list("rag_store_cluster3.json") or load_json_list("rag_store_cluster3.json.gz")
    st.session_state.stores = {
        "primary": primary,     # magazines / first corpus
        "secondary": secondary  # Cluster 3 / second corpus
    }

# Optional status
st.caption(
    "Corpora loaded — "
    f"primary: {len(st.session_state.stores['primary'])} chunks, "
    f"secondary: {len(st.session_state.stores['secondary'])} chunks."
)

# Parse credentials on each run
root_creds, user_creds = parse_credentials()

# Authentication
if not st.session_state.auth["is_authed"]:
    st.title("🔐 Sign In")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Log in"):
        if username in root_creds and root_creds[username] == password:
            st.session_state.auth = {"is_authed": True, "username": username, "role": "Root"}
            st.rerun()
        elif username in user_creds and user_creds[username][0] == password:
            role_name = user_creds[username][1]
            st.session_state.auth = {"is_authed": True, "username": username, "role": role_name}
            st.rerun()
        else:
            st.error("Invalid username or password")
    st.stop()

# ------------------------
# UI controls (top-level, not sidebar)
# ------------------------
st.info(f"Logged in as **{st.session_state.auth['username']}** ({st.session_state.auth['role']})")
if st.button("Log out"):
    st.session_state.auth = {"is_authed": False, "username": None, "role": None}
    st.session_state.messages = []
    st.session_state.docs = []
    st.session_state.all_chunks = []
    st.rerun()

colA, colB = st.columns([1, 6])
with colA:
    if st.button("Clear chat", key="btn_clear_chat"):
        st.session_state.messages = []
        st.rerun()
with colB:
    st.toggle("Stream responses (visual only)", value=True, key="stream_vis")

st.markdown("### 📎 Add documents")
uploads = st.file_uploader(
    "Drop files here (pdf, csv, xlsx, xls, docx, txt) — multiple allowed",
    type=["pdf", "csv", "xlsx", "xls", "docx", "txt"],
    accept_multiple_files=True,
    key="uploader_docs",
)

if st.button("Clear documents", key="btn_clear_docs_top"):
    st.session_state.docs = []
    st.session_state.all_chunks = []
    st.rerun()

# Process newly uploaded files
if uploads:
    max_docs = 2
    # Respect a maximum number of simultaneously loaded documents.
    # Skip any files that would exceed the limit and notify the user.
    for f in uploads:
        # If we've already reached the document limit, warn and stop processing further uploads
        if len(st.session_state.docs) >= max_docs:
            st.warning(f"Maximum of {max_docs} documents can be uploaded at once. Additional files were skipped.")
            break
        text, info, meta = extract_text_from_upload(f)
        if text:
            chunks = build_source_marked_chunks(f.name, meta, text)
            entry: Dict = {
                "name": f.name,
                "meta": info,
                "text": text,
                "chunks": chunks,
                "type": meta.get("type", "text"),
            }
            if meta.get("df") is not None:
                entry["df_preview"] = meta["df"].head(50)
            st.session_state.docs.append(entry)
            st.session_state.all_chunks.extend(chunks)
            st.success(f"Added: {info} (chunks: {len(chunks)})")
        else:
            st.warning(f"Skipped: {info}")

# Display loaded documents
if st.session_state.docs:
    st.markdown("#### 📚 Loaded documents")
    for i, d in enumerate(st.session_state.docs):
        with st.expander(f"{d['name']} — {d['meta']}", expanded=False):
            if d.get("df_preview") is not None:
                st.markdown("**Table preview (first 50 rows):**")
                st.dataframe(
                    d["df_preview"],
                    use_container_width=True,
                    hide_index=True,
                    key=f"dfprev_{i}",
                )
            st.text_area(
                "Extracted text (preview)",
                d["text"][:5000],
                height=180,
                key=f"preview_{i}_{d['name']}",
            )
            st.caption(
                "Source markers are added to retrieved snippets, e.g., [filename | type | page N] or [filename | type | rows i–j]."
            )

# Optional summarization
if st.session_state.docs:
    # Create a placeholder to manage summary output. This allows overwriting
    # the previous summary when generating a new one.
    if "summary_placeholder" not in st.session_state:
        st.session_state.summary_placeholder = st.empty()
    if st.button("🧠 Summarize loaded documents", key="summarize_btn"):
        context: List[str] = []
        for d in st.session_state.docs:
            sample = d["text"][:1500]
            meta_info = d["meta"]
            context.append(f"{d['name']} — {meta_info}\n{sample}")
        # Generate the summary text either using the OpenAI API or a local heuristic.
        summary_md = ""
        if HAS_OPENAI and os.getenv("OPENAI_API_KEY"):
            try:
                from openai import OpenAI  # type: ignore
                client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
                prompt = (
                    "Summarize the following documents in bullet points. Keep it concise and group by filename. "
                    "Quote key lines with [source markers].\n\n"
                    + "\n\n---\n\n".join(context)
                )
                resp = client.chat.completions.create(
                    model=os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini"),
                    messages=[
                        {"role": "system", "content": "You are a concise technical summarizer."},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.2,
                )
                summary_md = resp.choices[0].message.content or "(no content)"
            except Exception as e:
                st.error(f"Model summary failed, falling back to local: {e}")
        # Fall back to a local summary if no OpenAI API or on error.
        if not summary_md:
            agg = "### Local summary\n"
            for d in st.session_state.docs:
                toks = tokenize(d["text"])
                common_terms = ", ".join(w for w, _ in Counter(toks).most_common(10))
                agg += f"- **{d['name']}**: {len(d['text'])} chars; top terms: {common_terms}\n"
            summary_md = agg
        # Clear any previous summary and render the new one.
        placeholder = st.session_state.summary_placeholder
        placeholder.empty()
        placeholder.markdown(summary_md)

# Separator
st.markdown("---")

# Render chat history
for idx, msg in enumerate(st.session_state.messages):
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# Chat input with retrieval augmentation
prompt = st.chat_input("Ask a question about the uploaded docs—or chat normally…", key="chat_input")
if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    # PRIORITY SEARCH: uploads → primary (rag_store) → secondary (rag_store_cluster3)
    K = 5
    ctx: List[str] = []

    upload_chunks = st.session_state.all_chunks or []
    if upload_chunks:
        # When multiple documents are uploaded, ensure we retrieve context from each document.
        # Group the chunks by the originating document name, extracted from the leading bracket.
        doc_chunks_map: Dict[str, List[str]] = {}
        for c in upload_chunks:
            # Each chunk begins with a marker like "[filename | type | ...]".
            # Split on the '|' character and strip the leading '[' to isolate the filename.
            parts = c.split("|")
            if parts:
                name_part = parts[0].strip()
                if name_part.startswith("["):
                    name_part = name_part[1:]
                doc_name = name_part.strip()
                doc_chunks_map.setdefault(doc_name, []).append(c)
        # Collect top chunks from each document separately to ensure both documents contribute context.
        for doc_name, chunks_list in doc_chunks_map.items():
            ctx.extend(top_chunks(prompt, chunks_list, k=2))
        # Limit total upload context to K to avoid overfilling.
        ctx = ctx[:K]

    need = K - len(ctx)
    if need > 0:
        primary = st.session_state.stores.get("primary", [])
        if primary:
            ctx.extend(top_chunks(prompt, primary, k=need))

    need = K - len(ctx)
    if need > 0:
        secondary = st.session_state.stores.get("secondary", [])
        if secondary:
            ctx.extend(top_chunks(prompt, secondary, k=need))

    with st.chat_message("assistant"):
        with st.spinner("Analyzing…"):
            reply = minimal_model_reply(
                prompt,
                st.session_state.messages,
                ctx,
                role_name=st.session_state.auth.get("role"),
            )
            reply = linkify_labels(reply)
            st.markdown(reply)

    st.session_state.messages.append({"role": "assistant", "content": reply})
# Footer tip
st.caption("Tip: install optional extras for richer parsing: `pypdf`, `python-docx`, `openpyxl`.")
