#!/usr/bin/env python3
# make_mag_rag_parallel.py
import argparse, json, re, sys, os
from pathlib import Path
from typing import Any, Dict, List, Tuple
from concurrent.futures import ProcessPoolExecutor, as_completed
import multiprocessing

# ---------- feature toggles / deps ----------
def _has_pkg(name: str) -> bool:
    try: __import__(name); return True
    except Exception: return False
HAS_PYPDF = _has_pkg("pypdf")
HAS_DOCX  = _has_pkg("docx")  # <— add this

# ---------- small text utils ----------
import re
_NONWS = re.compile(r"\S")
_LINEBREAKS = re.compile(r"\r\n|\r|\n")
_EOL_HYPHEN = re.compile(r"(\w)-\n(\w)")
def normalize_ws(s: str) -> str:
    s = s.replace("\x00","")
    s = _LINEBREAKS.sub("\n", s)
    s = re.sub(r"[ \t]+\n", "\n", s)
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()
def fix_hyphenation(s: str) -> str:
    while True:
        new = _EOL_HYPHEN.sub(r"\1\2", s)
        if new == s: return new
        s = new
def chunk_text(s: str, chunk_chars: int = 1400, overlap: int = 200) -> List[str]:
    s = normalize_ws(s)
    if not s: return []
    if len(s) <= chunk_chars: return [s]
    out, start, L = [], 0, len(s)
    while start < L:
        end = min(L, start + chunk_chars)
        seg = s[start:end]
        if seg: out.append(seg)
        if end == L: break
        start = max(0, end - overlap)
    return out

# ---------- boilerplate (header/footer) ----------
from collections import Counter
def detect_boilerplate(lines_per_page: List[List[str]],
                       head_take: int = 3, tail_take: int = 3,
                       min_share: float = 0.2, top_n: int = 5) -> Tuple[set, set]:
    headers, footers, total = Counter(), Counter(), len(lines_per_page) or 1
    for lines in lines_per_page:
        if not lines: continue
        h = [l.strip() for l in lines[:head_take] if _NONWS.search(l or "")]
        f = [l.strip() for l in lines[-tail_take:] if _NONWS.search(l or "")]
        headers.update(h); footers.update(f)
    thresh = max(2, int(min_share * total))
    H = {t for t,c in headers.most_common() if c >= thresh}
    F = {t for t,c in footers.most_common() if c >= thresh}
    return set(list(H)[:top_n]), set(list(F)[:top_n])
def strip_boilerplate(page_text: str, header_lines: set, footer_lines: set) -> str:
    lines = page_text.split("\n")
    while lines and lines[0].strip() in header_lines: lines.pop(0)
    while lines and lines[-1].strip() in footer_lines: lines.pop()
    return "\n".join(lines)

# ---------- hash / dedup ----------
def sha1(s: str) -> str:
    import hashlib
    return hashlib.sha1(s.encode("utf-8")).hexdigest()
def dedup(chunks: List[str]) -> List[str]:
    seen, out = set(), []
    for c in chunks:
        h = sha1(c)
        if h in seen: continue
        seen.add(h); out.append(c)
    return out

# ---------- extraction (per file) ----------
def extract_pdf_text(path: Path) -> Tuple[str, Dict[str, Any], List[str]]:
    if not HAS_PYPDF:
        return "", {"type":"pdf_error","info":f"{path.name}: pypdf not installed"}, []
    from pypdf import PdfReader  # type: ignore
    try:
        reader = PdfReader(str(path))
    except Exception as e:
        return "", {"type":"pdf_error","info":f"{path.name}: {e}"}, []
    pages, empty_pages, chars_per_page = [], [], []
    for i, page in enumerate(reader.pages, start=1):
        try: txt = page.extract_text() or ""
        except Exception: txt = ""
        txt = fix_hyphenation(normalize_ws(txt))
        pages.append(txt)
        if not _NONWS.search(txt or ""): empty_pages.append(i)
        chars_per_page.append(len(txt))
    joined = "\n".join(pages)
    meta = {
        "type": "pdf",
        "info": f"{path.name} (PDF pages: {len(pages)})",
        "pages_count": len(pages),
        "empty_pages": empty_pages,
        "chars_per_page": chars_per_page,
    }
    return joined, meta, pages

def extract_docx_text(path: Path):
    """Return (joined_text, meta, 'pages') for DOCX (pages=[] but keep shape)."""
    if not HAS_DOCX:
        return "", {"type": "docx_error", "info": f"{path.name}: python-docx not installed"}, []
    from docx import Document  # type: ignore
    try:
        doc = Document(str(path))
    except Exception as e:
        return "", {"type": "docx_error", "info": f"{path.name}: {e}"}, []
    text = "\n".join(p.text for p in doc.paragraphs)
    text = fix_hyphenation(normalize_ws(text))
    meta = {
        "type": "docx",
        "info": f"{path.name} (DOCX paragraphs: {len(doc.paragraphs)})",
        "pages_count": None,
        "empty_pages": [],
        "chars_per_page": None,
    }
    # treat as single “page” for marker consistency
    return text, meta, [text]

def build_pdf_chunks(name: str, pages: List[str], chunk_chars: int,
                     overlap: int, do_strip_boilerplate: bool) -> Tuple[List[str], Dict[str, Any]]:
    audit: Dict[str, Any] = {}
    lines_per_page = [p.split("\n") for p in pages]
    header_lines, footer_lines = (set(), set())
    if do_strip_boilerplate and pages:
        header_lines, footer_lines = detect_boilerplate(lines_per_page)
        audit["boilerplate_header_lines"] = sorted(header_lines)
        audit["boilerplate_footer_lines"] = sorted(footer_lines)
    out: List[str] = []
    for i, page in enumerate(pages, start=1):
        if header_lines or footer_lines:
            page = strip_boilerplate(page, header_lines, footer_lines)
        for c in chunk_text(page, chunk_chars=chunk_chars, overlap=overlap):
            if c.strip():
                out.append(f"[{name} | pdf | page {i}] \n{c}")
    return out, audit

# ---------- worker (runs in a separate process) ----------
def _worker_parse(args) -> Tuple[str, Dict[str, Any], List[str]]:
    (root_str, path_str, chunk_chars, overlap, strip_bp) = args
    root, p = Path(root_str), Path(path_str)
    rel_path = str(p.relative_to(root))
    ext = p.suffix.lower()

    if ext == ".pdf":
        text, meta, pages = extract_pdf_text(p)
    elif ext == ".docx":
        text, meta, pages = extract_docx_text(p)
    else:
        # unsupported here
        return rel_path, {
            "name": p.name, "rel_path": rel_path, "type": "unsupported",
            "info": f"{p.name}: unsupported {ext}"
        }, []

    entry = {
        "name": p.name,
        "rel_path": rel_path,
        "type": meta.get("type"),
        "info": meta.get("info"),
        "pages_count": meta.get("pages_count"),
        "empty_pages": meta.get("empty_pages", []),
        "chunks_before_dedup": 0,
        "chunks_after_dedup": 0,
        "boilerplate_header_lines": [],
        "boilerplate_footer_lines": [],
    }
    if not text:
        return rel_path, entry, []

    # build chunks
    if ext == ".pdf":
        chunks, audit = build_pdf_chunks(
            name=p.name, pages=pages,
            chunk_chars=chunk_chars, overlap=overlap,
            do_strip_boilerplate=strip_bp
        )
        for k in ("boilerplate_header_lines","boilerplate_footer_lines"):
            if k in audit: entry[k] = audit[k]
    else:
        # DOCX: single “page” marker
        chunks = []
        for c in chunk_text(pages[0], chunk_chars=chunk_chars, overlap=overlap):
            if c.strip():
                chunks.append(f"[{p.name} | docx] \n{c}")

    entry["chunks_before_dedup"] = len(chunks)
    chunks = dedup(chunks)
    entry["chunks_after_dedup"] = len(chunks)
    return rel_path, entry, chunks
# ---------- main ----------
def main():
    ap = argparse.ArgumentParser(description="Parallel RAG builder for magazine PDFs.")
    ap.add_argument("-i","--input", required=True, help="Root input directory (recursed)")
    ap.add_argument("-o","--output", default="rag_store.json", help="Output JSON for chunks")
    ap.add_argument("--manifest", default="rag_manifest.json", help="Audit manifest JSON")
    ap.add_argument("--extensions", default="pdf", help="Comma-separated extensions (default: pdf)")
    ap.add_argument("--name-filter", default="", help="Regex to include only filenames that match (case-insensitive)")
    ap.add_argument("--chunk-chars", type=int, default=1400, help="Chars per text chunk")
    ap.add_argument("--overlap", type=int, default=200, help="Overlap chars between chunks")
    ap.add_argument("--strip-boilerplate", action="store_true", help="Remove repeating headers/footers")
    ap.add_argument("--workers", type=int, default=max(1, multiprocessing.cpu_count()-1),
                    help="Parallel workers (processes). Default: CPU count - 1")
    args = ap.parse_args()

    root = Path(args.input)
    if not root.exists() or not root.is_dir():
        print(f"[!] Input directory not found: {root}", file=sys.stderr)
        sys.exit(1)

    exts = {e.strip().lower().lstrip(".") for e in args.extensions.split(",") if e.strip()}
    name_rx = re.compile(args.name_filter, re.IGNORECASE) if args.name_filter else None

    files: List[Path] = []
    for p in root.rglob("*"):
        if not p.is_file(): continue
        if p.suffix.lower().lstrip(".") not in exts: continue
        if name_rx and not name_rx.search(p.name): continue
        files.append(p)
    files.sort(key=lambda p: str(p).lower())
    print(f"Discovered {len(files)} file(s). Using {args.workers} worker(s).")

    manifest: Dict[str, Any] = {
        "input_dir": str(root),
        "settings": {
            "extensions": sorted(list(exts)),
            "name_filter": args.name_filter or None,
            "chunk_chars": args.chunk_chars,
            "overlap": args.overlap,
            "strip_boilerplate": bool(args.strip_boilerplate),
            "workers": args.workers,
            "deps": { "pypdf": HAS_PYPDF },
        },
        "files": [],
        "summary": {},
    }

    # Submit tasks
    tasks = []
    with ProcessPoolExecutor(max_workers=args.workers) as ex:
        for p in files:
            tasks.append(ex.submit(
                _worker_parse,
                (str(root), str(p), args.chunk_chars, args.overlap, args.strip_boilerplate)
            ))

        # Collect results (progress prints in main process)
        results: List[Tuple[str, Dict[str,Any], List[str]]] = []
        for fut in as_completed(tasks):
            try:
                rel_path, entry, chunks = fut.result()
                if entry["chunks_after_dedup"] > 0:
                    print(f"[+] {Path(rel_path).name}: {entry['chunks_after_dedup']} chunks "
                          f"(pages={entry.get('pages_count','-')} empty={len(entry.get('empty_pages') or [])})")
                else:
                    print(f"[-] Skipped: {rel_path} — {entry.get('type')} ({entry.get('info')})")
                results.append((rel_path, entry, chunks))
            except Exception as e:
                print(f"[!] Worker failed: {e}", file=sys.stderr)

    # Deterministic order in outputs
    results.sort(key=lambda r: r[0].lower())
    all_chunks: List[str] = []
    for _, entry, chunks in results:
        manifest["files"].append(entry)
        all_chunks.extend(chunks)

    # Global de-dup
    before = len(all_chunks)
    all_chunks = dedup(all_chunks)
    after = len(all_chunks)
    manifest["summary"] = {
        "files_count": len(results),
        "chunks_total_before_dedup": before,
        "chunks_total_after_dedup": after,
    }

    out_fp = Path(args.output)
    man_fp = Path(args.manifest)
    out_fp.parent.mkdir(parents=True, exist_ok=True)
    man_fp.parent.mkdir(parents=True, exist_ok=True)

    # Force UTF-8 writes (fixes cp1252 issues on Windows)
    out_fp.write_text(json.dumps(all_chunks, ensure_ascii=False, indent=2), encoding="utf-8")
    man_fp.write_text(json.dumps(manifest,   ensure_ascii=False, indent=2), encoding="utf-8")

    print("\nDone.")
    print(f"Chunks: {after} → {out_fp.resolve()}")
    print(f"Manifest: {man_fp.resolve()}")
    print("Commit these JSONs so the Streamlit app can load them at deploy.")

if __name__ == "__main__":
    # Windows needs the guard for multiprocessing (spawn)
    main()